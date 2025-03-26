import os
import random
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from io import BytesIO
from google import genai
from datetime import datetime
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PIL import Image, ImageDraw, ImageFont
import textwrap
import json
import time
from dotenv import load_dotenv

load_dotenv()



# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.FileHandler('email_generator.log'), logging.StreamHandler()]
)
# Initialize Gemini API client
client = genai.Client(api_key=os.getenv('GEMINI_API_KEY'))
model = 'gemini-2.0-flash'

# Data generators
BANKS = ["Citizens Bank", "Wells Fargo", "Bank of America", "JPMorgan Chase"]
CUSTOMER_NAMES = ["John Smith", "Alice Johnson", "Robert Brown", "Emily Davis", "Michael Wilson", "Sophia Martinez"]
EMAIL_DOMAINS = ["example.com", "mail.com", "bankmail.com"]


def generate_financial_data():
    """Generate realistic financial test data with extra fields and random customer details."""
    customer_name = random.choice(CUSTOMER_NAMES)
    customer_email = f"{customer_name.lower().replace(' ', '.')}@{random.choice(EMAIL_DOMAINS)}"
    customer_phone = f"+1-{random.randint(200, 999)}-{random.randint(100, 999)}-{random.randint(1000, 9999)}"
    return {
        'loan_number': f"LN-{random.randint(100000, 999999)}",
        'amount': f"${random.randint(1000, 500000):,}.{random.randint(0, 99):02d}",
        'date': datetime.now().strftime('%Y-%m-%d'),
        'transaction_id': f"TX-{random.randint(1000000, 9999999)}",
        'vendor': random.choice(['AlphaTech Solutions', 'Beta Systems', 'Gamma Corp']),
        'invoice_number': f"INV-{random.randint(10000, 99999)}",
        'project_code': f"PRJ-{random.randint(100, 999)}",
        'budget': f"${random.randint(5000, 100000):,}",
        'fee': f"${random.randint(50, 1000):,}",
        'payer': random.choice(['Global Corp Inc', 'United Holdings Ltd', 'Prime Ventures LLC']),
        'payee': random.choice(['Merchant Bank International', 'First Capital Trust', 'Metropolitan Financial']),
        'account_number': ''.join(str(random.randint(0, 9)) for _ in range(12)),
        'routing_number': ''.join(str(random.randint(0, 9)) for _ in range(9)),
        'swift_code': f"SWIFT{random.choice(['A', 'B', 'C'])}{random.randint(1000, 9999)}",
        'customer_name': customer_name,
        'customer_email': customer_email,
        'customer_phone': customer_phone
    }


def generate_email_content(request_type, sub_type, num_attachments):
    """
    Generate email chain content and attachment narratives.
    The Gemini prompt specifies the exact number of attachments and requires a JSON object with:
      - "email_chain": complete HTML email chain (descending order with proper <html><body> tags)
      - "attachments": an array (of length num_attachments) with distinct narrative texts for each attachment.
    """
    financial_data = generate_financial_data()
    prompt = f"""
I want you to generate commercial banking email content for testing an email classification system.

WORKFLOW:
The generated emails should represent a workflow where:
1. A customer sends a request to the bank's front office.
2. The front office collects documents and information.
3. The front office then forwards the request with the collected documents to the back office for processing.

CONTEXT:
The internal request type is "{request_type}" with sub-category "{sub_type}".
Do not mention these exact terms in the email chain content. Instead, subtly incorporate their essence.

EMAIL STRUCTURE:
- Write in a narrative, conversational style with natural language paragraphs.
- Include realistic financial details (dates, amounts, account numbers, transaction IDs, vendor names, etc.) naturally.
- Generate an authentic email chain (with replies and forwards) in descending order (newest first).
- Do not include any "Subject:" header lines within the email body.
- Clearly reference that supporting documents are attached (for example, "please see attached report 1", "attached report 2", etc.).
- Ensure that all the names are fakely generated and ensure that there are no placeholders, the email should look completely authentic

ATTACHMENT CONTENT:
- Generate narrative text for attachments that provide detailed context and financial reconciliation information.
- Produce exactly {num_attachments} distinct narrative texts. Each should have a different style.
- If the email chain has less details the attachment should be very detailed and should have loads of information, and if the email chain has more information, the attachment can have barebones details.
- Do not add any extra additional details.

OUTPUT:
Return ONLY a valid JSON object (with no extra text) with exactly two keys:
{{
  "email_chain": "<complete HTML email chain content>",
  "attachments": [ "<narrative text for attachment 1>", "<narrative text for attachment 2>", ... ]
}}
Ensure that the "attachments" array has exactly {num_attachments} elements and that the JSON does not include any markdown formatting.
    """
    logging.info(f"Generating email content for {request_type}/{sub_type} with {num_attachments} attachments")
    start_time = time.time()
    try:
        response = client.models.generate_content(model=model, contents=prompt)
        if not response.text:
            raise ValueError("Empty response from Gemini API")
        output = response.text.strip()
        if output.startswith("```"):
            first_newline = output.find("\n")
            if first_newline != -1:
                output = output[first_newline + 1:]
            if output.endswith("```"):
                output = output[:-3].strip()
        try:
            result = json.loads(output)
            if "attachments" not in result or len(result["attachments"]) != num_attachments:
                logging.error("Generated attachments array length does not match expected number.")
                raise ValueError("Incorrect attachments array length.")
        except Exception as e:
            logging.error("Failed to parse JSON response; response was:")
            logging.error(output)
            result = {
                "email_chain": f"<html><body>Error parsing email chain content for {request_type}/{sub_type}.</body></html>",
                "attachments": [f"Error generating attachment content for {request_type}/{sub_type}."] * num_attachments
            }
        logging.info(f"Generated email content in {time.time() - start_time:.2f}s")
        return result
    except Exception as e:
        logging.error(f"Content generation failed: {str(e)}")
        return {
            "email_chain": f"<html><body>Error generating email content for {request_type}/{sub_type}</body></html>",
            "attachments": [f"Error generating attachment content for {request_type}/{sub_type}."] * num_attachments
        }


def create_attachment_from_content(content, variant_suffix=""):
    """
    Generate an attachment file from the given narrative text.
    Randomly choose a file format (pdf, docx, or png) and return the attachment name and its binary data.
    The variant_suffix is appended to the content (if needed) to ensure differences.
    """
    base_data = generate_financial_data()
    modified_content = content + "\n" + variant_suffix
    file_choice = random.choice(['pdf', 'docx', 'png'])
    if file_choice == 'pdf':
        filename = f"report_{base_data['loan_number']}.pdf"
        data = create_pdf_attachment(modified_content)
    elif file_choice == 'docx':
        filename = f"report_{base_data['loan_number']}.docx"
        data = create_docx_attachment(modified_content)
    else:
        filename = f"report_{base_data['loan_number']}.png"
        data = create_image_attachment(modified_content)
    return filename, data


def create_image_attachment(content):
    """Generate a PNG attachment with dynamic height, colourful drawings, and a bank logo."""
    margin = 40
    line_height = 24
    lines = []
    for line in content.split('\n'):
        if line.strip():
            wrapped = textwrap.wrap(line, width=60)
            lines.extend(wrapped)
        else:
            lines.append("")
    num_lines = len(lines)
    height = margin * 2 + num_lines * line_height
    width = 800
    image = Image.new('RGBA', (width, height), (255, 255, 255, 255))
    draw = ImageDraw.Draw(image)
    draw.rectangle([(0, 0), (width - 1, height - 1)], outline=(0, 128, 255), width=3)
    draw.ellipse([(width - 120, 10), (width - 20, 110)], outline=(255, 0, 0), width=3)
    draw.line([(0, height - 30), (width, height - 30)], fill=(0, 255, 0), width=4)
    bank = random.choice(BANKS)
    logo_width, logo_height = 150, 50
    draw.rectangle([(10, 10), (10 + logo_width, 10 + logo_height)], fill=(200, 200, 200), outline=(0, 0, 0), width=2)
    try:
        logo_font = ImageFont.truetype("Arial", 18)
    except IOError:
        logo_font = ImageFont.load_default()
    draw.text((20, 20), bank, font=logo_font, fill=(0, 0, 128))
    try:
        font = ImageFont.truetype("Arial", 16)
    except IOError:
        font = ImageFont.load_default()
    y_position = margin
    for line in lines:
        draw.text((margin, y_position), line, font=font, fill=(0, 0, 0))
        y_position += line_height
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer.getvalue()


def create_pdf_attachment(content):
    """Generate a PDF attachment with proper pagination so that text does not overflow."""
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    margin = 40
    line_height = 14
    max_lines = int((height - 2 * margin) / line_height)
    wrapped_lines = []
    for line in content.split('\n'):
        if line.strip():
            wrapped = textwrap.wrap(line, width=90)
            wrapped_lines.extend(wrapped)
        else:
            wrapped_lines.append("")
    current_line = 0
    while current_line < len(wrapped_lines):
        text_obj = c.beginText(margin, height - margin)
        text_obj.setFont("Helvetica", 12)
        for _ in range(max_lines):
            if current_line >= len(wrapped_lines):
                break
            text_obj.textLine(wrapped_lines[current_line])
            current_line += 1
        c.drawText(text_obj)
        if current_line < len(wrapped_lines):
            c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.getvalue()


def create_docx_attachment(content):
    """Generate a DOCX attachment with the given narrative text."""
    doc = Document()
    doc.add_heading('Detailed Narrative Report', 0)
    for line in content.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def enforce_html_structure(html_content):
    """Ensure HTML content is wrapped with <html> and <body> tags."""
    if not html_content.strip().lower().startswith("<html"):
        html_content = f"<html><body>{html_content}</body></html>"
    return html_content


def build_email(request_type, sub_type):
    """
    Construct the complete email.
    Calls generate_email_content() with the specified number of attachments,
    builds the email chain, creates attachments (each from a distinct narrative),
    and appends attachment filename references to the email chain.
    Returns the complete MIME message.
    """
    num_attachments = random.randint(1, 3)
    content_result = generate_email_content(request_type, sub_type, num_attachments)
    email_html = content_result.get("email_chain",
                                    f"<html><body>No content generated for {request_type}/{sub_type}.</body></html>")
    attachments_array = content_result.get("attachments", [])
    if len(attachments_array) != num_attachments:
        logging.error("Attachments array length mismatch; expected %d, got %d.", num_attachments,
                      len(attachments_array))
        return None

    # Remove any "Subject:" lines from the email chain.
    lines = email_html.splitlines()
    filtered_lines = [line for line in lines if not line.strip().startswith("Subject:")]
    email_html = "\n".join(filtered_lines)
    email_html = enforce_html_structure(email_html)

    # Create attachments for each narrative.
    attachment_files = []
    for attachment_text in attachments_array:
        filename, data = create_attachment_from_content(attachment_text)
        attachment_files.append((filename, data))

    # Append references to the attachments at the end of the email chain.
    attachment_ref_html = "<br><br><p>Attached Documents:</p><ul>"
    for fname, _ in attachment_files:
        attachment_ref_html += f"<li>{fname}</li>"
    attachment_ref_html += "</ul>"
    email_html += attachment_ref_html

    # Create MIME message with a generic subject.
    financial_data = generate_financial_data()
    msg = MIMEMultipart('mixed')
    msg['From'] = f"Front Office <frontoffice@{random.choice(BANKS).lower().replace(' ', '')}.com>"
    msg['To'] = f"Back Office <backoffice@{random.choice(BANKS).lower().replace(' ', '')}.com>"
    msg['Subject'] = f"Banking Communication Ref: {financial_data['loan_number']}"
    msg['Date'] = datetime.now().strftime('%a, %d %b %Y %H:%M:%S %z')

    alt_part = MIMEMultipart('alternative')
    plain_text = "This email contains HTML content. Please use an HTML-compatible email client to view it properly."
    alt_part.attach(MIMEText(plain_text, 'plain'))
    alt_part.attach(MIMEText(email_html, 'html'))
    msg.attach(alt_part)

    # Attach each file with correct MIME headers and base64 encoding.
    for filename, data in attachment_files:
        if filename.endswith('.pdf'):
            content_type = 'application/pdf'
        elif filename.endswith('.docx'):
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif filename.endswith('.png'):
            content_type = 'image/png'
        else:
            content_type = 'application/octet-stream'
        attachment_part = MIMEBase(*content_type.split('/', 1))
        attachment_part.set_payload(data)
        encoders.encode_base64(attachment_part)
        attachment_part.add_header('Content-Disposition', 'attachment', filename=filename)
        attachment_part.add_header('Content-Type', content_type, name=filename)
        msg.attach(attachment_part)

    return msg


def save_email(email_msg, req_type, sub_type):
    """Save email as .eml under folder structure <output_dir>/<request_type>/<sub_request_type>/<email_number>.eml."""
    email_number = f"{random.randint(100, 999)}.eml"
    output_dir = "generated_emails"
    req_dir = os.path.join(output_dir, req_type.replace(" ", "_"))
    sub_req_dir = os.path.join(req_dir, sub_type.replace(" ", "_"))
    os.makedirs(sub_req_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = os.path.join(sub_req_dir, f"{timestamp}_{email_number}")
    with open(filename, 'w') as f:
        f.write(email_msg.as_string())
    logging.info(f"Saved email: {filename}")


# Main execution
if __name__ == "__main__":
    # Expanded request types.
    REQUEST_TYPES = [
        ("adjustment", ["reallocation fees", "amendment fees", "reallocation principal"]),
        ("au transfer", ["domestic", "international"]),
        ("closing notice", ["pre-closing", "post-closing"]),
        ("commitment change", ["cashless roll", "decrease", "increase"]),
        ("fee payment", ["ongoing fee", "letter of credit fee"]),
        ("money movement inbound", ["principal", "interest", "principal + interest", "principal + interest + fee"]),
        ("money movement outbound", ["timebound", "foreign currency"]),
        ("document request", ["identity verification", "income statement"]),
        ("account update", ["contact details", "address change"]),
        ("loan extension", ["term extension", "rate modification"]),
        ("collateral review", ["valuation update", "insurance verification"]),
        ("regulatory compliance", ["audit request", "report submission"]),
        ("portfolio rebalancing", ["asset reallocation", "liquidity adjustment"]),
        ("risk assessment", ["credit risk", "market risk"]),
        ("operational review", ["system update", "process improvement"]),
        ("cash management", ["daily reconciliation", "overdraft monitoring"]),
        ("credit line adjustment", ["increase limit", "decrease limit"]),
        ("customer onboarding", ["identity check", "document submission"]),
        ("investment advisory", ["asset allocation", "risk profiling"])
    ]

    # Generate 3 emails for each request/sub-request type.
    for req_type, sub_types in REQUEST_TYPES:
        sub_types = sub_types or [""]
        for sub_type in sub_types:
            for _ in range(3):
                try:
                    email_msg = build_email(req_type, sub_type)
                    if email_msg is not None and "Error generating email content" not in email_msg.as_string():
                        save_email(email_msg, req_type, sub_type)
                    else:
                        logging.error(f"Email generation error for {req_type}/{sub_type} - skipping save.")
                except Exception as e:
                    logging.error(f"Failed processing {req_type}/{sub_type}: {str(e)}")
    logging.info("Process completed. Emails generated in the folder structure 'generated_emails'.")
